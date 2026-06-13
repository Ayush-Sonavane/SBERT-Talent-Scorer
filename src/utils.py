try:
    import orjson
    _loads = orjson.loads
except ImportError:
    import json
    _loads = json.loads

import docx


def load_docx_text(file_path: str) -> str:
    doc = docx.Document(file_path)
    texts = []

    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            seen_cells = set()
            row_text = []
            for cell in row.cells:
                if cell not in seen_cells:
                    seen_cells.add(cell)
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
            if row_text:
                texts.append(" | ".join(row_text))

    return "\n".join(texts)


def stream_jsonl(file_path: str):
    with open(file_path, "r", encoding="utf-8") as f:
        for idx, line in enumerate(f, 1):
            line = line.strip()
            if line:
                try:
                    yield _loads(line)
                except Exception as e:
                    print(f"Warning: Skipping malformed JSON line at row {idx}: {e}")
                    continue
